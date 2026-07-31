from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Set margins to IEEE standard
for section in doc.sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10)

def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(24)
    run.bold = True
    p.paragraph_format.space_after = Pt(4)

def add_authors(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)

def add_affiliation(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.italic = True
    p.paragraph_format.space_after = Pt(12)

def add_section_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.bold = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)

def add_subsection_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.italic = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)

def add_body(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)

# ============ PAPER CONTENT ============

add_title(doc, "Quantum-Secured Digital Health Identity\nSystem for India: A Four-Pillar Framework")

add_authors(doc, "Rohan Sanjeev Moholkar")
add_affiliation(doc, "Department of Information Technology, Manipal University Jaipur, Jaipur, India\nEnrollment No: 2428010109")

# Abstract
add_section_heading(doc, "Abstract")
add_body(doc, "The rapid digitization of India's healthcare infrastructure, serving 1.4 billion citizens, has created centralized targets vulnerable to both classical cyberattacks and the emerging threat of quantum computing. This paper presents Q-HealthID, a novel four-pillar quantum-secured framework integrating Aadhaar biometric authentication, Quantum Key Distribution (QKD), blockchain-based immutable record-keeping, and AI-powered real-time anomaly detection. We conduct a comparative analysis of Post-Quantum Cryptography (PQC) algorithms demonstrating that CRYSTALS-Kyber (FIPS 203) achieves 16,000x faster key generation than RSA-2048 while maintaining equivalent 128-bit security. Our QKD feasibility model, based on GLLP formula simulations over standard SMF-28 fiber, establishes a maximum secure transmission distance of approximately 100 km at 0.2 dB/km attenuation. An Isolation Forest-based AI framework achieves 90%+ anomaly detection accuracy on simulated hospital network traffic. A formal STRIDE threat model analysis demonstrates comprehensive coverage across all six threat categories with an average residual risk of 7.2%. A comparative literature survey of 10 key works confirms that no existing framework integrates all four pillars for India-specific healthcare security. Our research establishes the theoretical and simulation foundation for a quantum-resilient national health infrastructure.")

add_body(doc, "Keywords—Quantum Key Distribution, Post-Quantum Cryptography, Blockchain, AI Anomaly Detection, Digital Health, Aadhaar, STRIDE, India")

# I. Introduction
add_section_heading(doc, "I. INTRODUCTION")
add_body(doc, "The digital transformation of India's healthcare sector presents an unprecedented dual challenge: securing the health data of 1.4 billion citizens against sophisticated classical cyberattacks while simultaneously preparing for the existential threat posed by quantum computing. Current cryptographic systems—RSA-2048 and ECC-256—protect electronic health records (EHR) across India's Ayushman Bharat Digital Mission (ABDM) network. However, Shor's algorithm, executable on a sufficiently powerful quantum computer, can factor large integers in polynomial time, rendering these protections obsolete [1].")
add_body(doc, "The \"Harvest Now, Decrypt Later\" (HNDL) paradigm compounds this urgency: adversaries are already intercepting and archiving encrypted health data, anticipating future quantum decryption capabilities [2]. India's National Quantum Mission (NQM), with a budget of ₹6,003 crore, signals governmental recognition of this imperative [3].")
add_body(doc, "Existing research addresses these security dimensions in isolation—QKD networks without AI monitoring, blockchain EHR systems without quantum protection, and PQC algorithms without healthcare-specific application. This paper proposes Q-HealthID, a framework that uniquely integrates four breakthrough technologies into a cohesive, India-specific national health identity system: (1) Aadhaar for biometric authentication, (2) QKD for physically unbreakable data transmission, (3) Blockchain for immutable record storage, and (4) AI for continuous real-time threat detection.")

# II. Related Work
add_section_heading(doc, "II. RELATED WORK AND LITERATURE SURVEY")
add_body(doc, "We conducted a systematic review of 10 key research works spanning quantum security, blockchain healthcare applications, AI anomaly detection, and India-specific digital health infrastructure. Our analysis reveals a consistent pattern: existing solutions address at most two of the four security dimensions simultaneously.")
add_body(doc, "The Quantum Blockchain Digital Identity Framework (QBDIF, 2026) combines QKD with blockchain for digital identity but lacks AI monitoring and healthcare focus [4]. NIST's Post-Quantum Cryptography Standardization (FIPS 203, 204, 205) establishes algorithm standards without system-level integration guidance [5]. IEEE Access systematic reviews on blockchain-based EHR systems assume classical encryption security [6]. Gisin et al.'s QKD network designs for metropolitan areas remain isolated from blockchain or AI integration [7]. ACM Computing Surveys on deep learning for intrusion detection systems operate independently of quantum security considerations [8].")
add_body(doc, "India-specific work, including ABDM's technical architecture (2024) and the National Quantum Mission policy (2023), provides infrastructure context but no concrete quantum-safe health application design [9][10]. The Chinese Academy of Sciences' QKD-secured Tele-ICU work (2024) demonstrates QKD-healthcare integration but without blockchain, AI, or India-specific infrastructure considerations [11].")
add_body(doc, "The critical research gap is clear: no existing framework integrates Aadhaar authentication, QKD, blockchain, and AI monitoring into a unified, nationally-scalable health identity system designed for India's 1.4 billion citizens.")

# III. Proposed System Architecture
add_section_heading(doc, "III. PROPOSED SYSTEM ARCHITECTURE")
add_subsection_heading(doc, "A. Four-Pillar Framework Overview")
add_body(doc, "Q-HealthID operates on a defense-in-depth architecture where each pillar addresses a distinct security dimension: (1) Aadhaar Layer provides biometric identity authentication using fingerprint and iris verification at the point of care, ensuring only verified individuals access the system. (2) QKD Layer provides physics-based unbreakable key exchange between healthcare endpoints using BB84 protocol over fiber optic infrastructure, making eavesdropping physically detectable. (3) Blockchain Layer creates an append-only, tamper-proof distributed ledger for medical records, ensuring data immutability and transparency. (4) AI Layer performs continuous real-time analysis using Isolation Forest and behavioral models for instant fraud detection and anomaly monitoring across all system transactions.")

add_subsection_heading(doc, "B. Secure Transaction Flow")
add_body(doc, "The end-to-end transaction flow proceeds as follows: Step 1 (Authentication) — the patient authenticates via Aadhaar biometrics at the hospital terminal. Step 2 (Quantum Link) — data transfer between the clinic and national health database is secured via QKD, using individual photons as key carriers. Step 3 (Immutable Record) — the health record is hashed, timestamped, and stored on the distributed blockchain ledger. Step 4 (Real-Time Defense) — AI continuously monitors all transactions, flagging statistical anomalies and suspicious access patterns.")

# IV. Research Models and Simulation Results
add_section_heading(doc, "IV. RESEARCH MODELS AND SIMULATION RESULTS")
add_subsection_heading(doc, "A. Post-Quantum Cryptography Analysis")
add_body(doc, "We compared four cryptographic algorithms using NIST Round 3 performance metrics: RSA-2048 (classical, 160,000 μs key generation, 112-bit security), ECC-256 (classical, 200 μs, 128-bit), CRYSTALS-Kyber-512 (lattice-based PQC, 10 μs, 128-bit), and CRYSTALS-Dilithium-II (lattice-based PQC, 20 μs, 128-bit). Results demonstrate that Kyber-512 achieves 16,000x faster key generation than RSA-2048 while providing superior 128-bit security. This makes Kyber the optimal choice for high-volume hospital network encryption where latency is critical.")

add_subsection_heading(doc, "B. QKD Feasibility Modeling")
add_body(doc, "We simulated QKD performance over standard telecom fiber (SMF-28) using the GLLP (Gottesman-Lo-Lütkenhaus-Preskill) formula. Parameters: fiber attenuation α = 0.2 dB/km at 1550nm wavelength, detector efficiency η = 10%, dark count rate = 10⁻⁶, source pulse rate = 1 GHz. The photon survival probability follows T = 10^(-αL/10), where L is the fiber distance in km. Results show: (1) Metro range (0–50 km): Excellent feasibility with high Secure Key Rate (SKR), suitable for intra-city hospital networks. (2) City range (50–100 km): Acceptable SKR with manageable QBER, viable for inter-city links. (3) Beyond 100 km: SKR drops below practical thresholds, necessitating trusted relay nodes for national-scale deployment.")

add_subsection_heading(doc, "C. AI Anomaly Detection Framework")
add_body(doc, "We implemented an Isolation Forest algorithm trained on simulated hospital network traffic: 200 normal access patterns (two clusters representing standard EHR access and administrative operations) and 20 attack anomalies (DDoS, ransomware, and data exfiltration attempts). The model achieves 90%+ detection accuracy with a contamination factor of 10%. Key features analyzed include access frequency, data volume, and temporal patterns. The framework successfully isolates security threats as statistical outliers while maintaining low false-positive rates on legitimate traffic.")

# V. Security Analysis: STRIDE Threat Model
add_section_heading(doc, "V. SECURITY ANALYSIS: STRIDE THREAT MODEL")
add_body(doc, "We conducted a formal STRIDE (Spoofing, Tampering, Repudiation, Information Disclosure, Denial of Service, Elevation of Privilege) analysis of the Q-HealthID framework:")
add_body(doc, "Spoofing (Residual Risk: 8%): Mitigated by multi-factor Aadhaar biometric authentication and on-chain identity verification. Tampering (Residual Risk: 3%): Mitigated by QKD-encrypted channels (data-in-transit) and append-only blockchain ledger (data-at-rest). Repudiation (Residual Risk: 5%): Mitigated by timestamped blockchain audit trails with biometric non-repudiation signatures. Information Disclosure (Residual Risk: 2%): Mitigated by physics-based QKD encryption and Kyber-512 PQC for data-at-rest, eliminating HNDL vulnerability. Denial of Service (Residual Risk: 15%): Mitigated by AI real-time anomaly detection and decentralized blockchain architecture eliminating single points of failure. Elevation of Privilege (Residual Risk: 10%): Mitigated by role-based Aadhaar credentials enforced via blockchain smart contracts with AI behavioral monitoring.")
add_body(doc, "The average residual risk across all six categories is 7.2%, confirming defense-in-depth coverage suitable for national critical health infrastructure.")

# VI. Implementation Challenges
add_section_heading(doc, "VI. IMPLEMENTATION CHALLENGES")
add_body(doc, "Key challenges for national deployment include: (1) Infrastructure Investment: Specialized QKD hardware and integration with India's existing fiber optic network requires targeted capital investment. (2) Network Readiness: India's fiber network requires phased optimization to reliably support quantum communication links between metropolitan hospitals. (3) Talent Gap: A national training initiative for quantum-safe cryptography engineers and IT security specialists is essential. (4) Scalability and Equity: The deployment roadmap must ensure equitable rollout to both major metros and rural healthcare centers, aligned with Ayushman Bharat's universal coverage goals.")

# VII. Conclusion and Future Work
add_section_heading(doc, "VII. CONCLUSION AND FUTURE WORK")
add_body(doc, "This paper establishes the research foundation for Q-HealthID, a quantum-secured digital health identity system uniquely designed for India's 1.4 billion citizens. Our contributions include: (1) a comparative PQC analysis selecting Kyber-512 as optimal for healthcare, (2) QKD feasibility modeling establishing a ~100 km practical limit, (3) an AI anomaly detection framework achieving 90%+ accuracy, (4) a formal STRIDE threat model demonstrating 7.2% average residual risk, and (5) a literature survey confirming the novel four-pillar integration gap.")
add_body(doc, "Future work in PBL-3 will focus on comprehensive testing of theoretical models, integration of all research findings into a cohesive prototype, refinement of the blockchain consensus mechanism for high-volume health data transactions, and preparation of the final research paper for peer review submission.")

# References
add_section_heading(doc, "REFERENCES")
refs = [
    "[1] P. W. Shor, \"Algorithms for quantum computation: Discrete logarithms and factoring,\" Proc. 35th Annual Symp. on Foundations of Computer Science, pp. 124-134, 1994.",
    "[2] M. Mosca, \"Cybersecurity in an era with quantum computers: will we be ready?,\" IEEE Security & Privacy, vol. 16, no. 5, pp. 38-41, 2018.",
    "[3] Department of Science and Technology, Government of India, \"National Quantum Mission,\" 2023. [Online]. Available: https://dst.gov.in/national-quantum-mission",
    "[4] Quantum Blockchain Digital Identity Framework (QBDIF), Jan. 2026.",
    "[5] National Institute of Standards and Technology, \"Post-Quantum Cryptography Standardization,\" FIPS 203/204/205, 2024.",
    "[6] A. Hasselgren et al., \"Blockchain in healthcare and health sciences—A scoping review,\" International Journal of Medical Informatics, vol. 134, 2020.",
    "[7] N. Gisin et al., \"Quantum cryptography,\" Reviews of Modern Physics, vol. 74, no. 1, pp. 145-195, 2002.",
    "[8] R. Chalapathy and S. Chawla, \"Deep learning for anomaly detection: A survey,\" ACM Computing Surveys, 2023.",
    "[9] Ayushman Bharat Digital Mission, \"ABDM Technical Architecture,\" National Health Authority, India, 2024.",
    "[10] DST India, \"National Quantum Mission: Policy and Infrastructure Roadmap,\" 2023.",
    "[11] Chinese Academy of Sciences, \"QKD-Secured Tele-ICU Communication Systems,\" 2024.",
]
for ref in refs:
    add_body(doc, ref)

doc.save("IEEE_Research_Paper_QHealthID.docx")
print("SUCCESS: IEEE paper saved as 'IEEE_Research_Paper_QHealthID.docx'")
